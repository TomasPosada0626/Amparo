"""Extraccion de pares (consulta, respuesta esperada) desde .md, .docx o .pdf.

El formato esperado es el mismo que private/dataset_legal_30_ejemplos.md:

    Entrada:
    "..."

    Salida esperada:
    "..."

repetido para cada ejemplo, opcionalmente precedido por un encabezado tipo
"### Ejemplo N". Los sinonimos Consulta/Pregunta y Respuesta esperada tambien
se reconocen para tolerar variaciones menores en archivos Word/PDF.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Optional


@dataclass
class DatasetItem:
    query: str
    expected: Optional[str] = None


class DatasetParseError(RuntimeError):
    pass


_ENTRY_RE = re.compile(
    r"^[ \t]*(?:Entrada|Consulta|Pregunta)\s*:\s*(?P<query>.*?)\s*"
    r"^[ \t]*(?:Salida|Respuesta)\s+Esperada\s*:\s*(?P<expected>.*?)"
    r"(?=^[ \t]*(?:#{0,6}\s*Ejemplo\b|(?:Entrada|Consulta|Pregunta)\s*:)|\Z)",
    re.IGNORECASE | re.DOTALL | re.MULTILINE,
)

_QUOTE_PAIRS = {'"': '"', "'": "'", "“": "”", "‘": "’"}


def _strip_quotes(text: str) -> str:
    text = text.strip()
    if len(text) >= 2 and text[0] in _QUOTE_PAIRS and text[-1] == _QUOTE_PAIRS[text[0]]:
        text = text[1:-1].strip()
    return text


def _extract_docx(path: Path) -> str:
    from docx import Document  # import perezoso: solo se necesita para .docx

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader  # import perezoso: solo se necesita para .pdf

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in (".md", ".txt"):
        return path.read_text(encoding="utf-8")
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    raise DatasetParseError(f"Formato no soportado: '{suffix}'. Usa .md, .docx o .pdf.")


def parse_dataset(path: Path) -> list[DatasetItem]:
    """Parsea un archivo con pares Entrada/Salida esperada.

    Lanza DatasetParseError si no se reconoce ningun par, para evitar lanzar
    una comparacion "vacia" en silencio cuando el archivo no sigue el formato.
    """
    text = _extract_text(path)
    items: list[DatasetItem] = []
    for match in _ENTRY_RE.finditer(text):
        query = _strip_quotes(match.group("query"))
        expected = _strip_quotes(match.group("expected"))
        if query:
            items.append(DatasetItem(query=query, expected=expected or None))

    if not items:
        raise DatasetParseError(
            f"No se encontraron pares 'Entrada: / Salida esperada:' en '{path.name}'. "
            "Verifica que el archivo siga el mismo formato que "
            "private/dataset_legal_30_ejemplos.md."
        )
    return items
